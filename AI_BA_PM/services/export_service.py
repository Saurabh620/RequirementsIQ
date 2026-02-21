"""
Service: Export Service
Generates professional PDF and DOCX files from document artifacts.
PDF: ReportLab | DOCX: python-docx
"""
import io
from datetime import datetime
from typing import Optional

# ReportLab imports at module level so helper functions can access them
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, HRFlowable, PageBreak)
except ImportError:
    pass  # Will fail gracefully inside generate_pdf if not installed


# ── PDF Export (ReportLab) ────────────────────────────────────
def generate_pdf(document_data: dict) -> bytes:
    """
    Render a professional PDF from the document data dict.
    Returns raw PDF bytes.
    """

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=20*mm, leftMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm
    )

    styles = getSampleStyleSheet()
    ACCENT = colors.HexColor("#6366f1")
    DARK = colors.HexColor("#1e1b4b")

    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                  textColor=DARK, fontSize=22, spaceAfter=4)
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"],
                               textColor=ACCENT, fontSize=14, spaceBefore=12, spaceAfter=4)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                               textColor=DARK, fontSize=12, spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                 fontSize=10, leading=14, spaceAfter=4)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.gray)

    story = []
    title = document_data.get("title", "Requirements Document")
    domain = document_data.get("domain", "generic").upper()
    created = document_data.get("created_at", datetime.now().strftime("%B %d, %Y"))

    # Header
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Domain: {domain} | Generated: {created} | Powered by RequirementIQ", meta_style))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT))
    story.append(Spacer(1, 8*mm))

    # ── BRD Section ──────────────────────────────────────────
    brd = document_data.get("brd")
    if brd:
        story.append(Paragraph("BUSINESS REQUIREMENTS DOCUMENT", h1_style))
        
        if brd.get("project_name"):
            story.append(Paragraph(f"Project: {brd['project_name']}", h2_style))

        if brd.get("document_control"):
            story.append(Paragraph("1. Document Control", h2_style))
            dc = brd["document_control"]
            dc_data = [["Field", "Details"]]
            for k, v in dc.items():
                dc_data.append([k.replace("_", " ").title(), str(v)])
            t = Table(dc_data, colWidths=[50*mm, 120*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("FONTSIZE", (0,0), (-1,-1), 9),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        _add_section(story, "2. Executive Summary", brd.get("executive_summary", {}).get("content"), h2_style, body_style)
        _add_list_section(story, "3. Business Objectives", brd.get("business_objectives", []), h2_style, body_style)
        _add_list_section(story, "Success Criteria", brd.get("success_criteria", []), h2_style, body_style)
        _add_section(story, "4. Problem Statement", brd.get("problem_statement", {}).get("content"), h2_style, body_style)
        _add_list_section(story, "5.1 Scope — In", brd.get("scope_in", []), h2_style, body_style)
        _add_list_section(story, "5.2 Scope — Out", brd.get("scope_out", []), h2_style, body_style)

        # 6. Stakeholders table
        stakeholders = [
            s for s in brd.get("stakeholders", [])
            if s.get("name","") not in ("", "INSUFFICIENT_DATA")
            and s.get("role","") not in ("", "INSUFFICIENT_DATA")
        ]
        if stakeholders:
            story.append(Paragraph("6. Stakeholders", h2_style))
            tdata = [["Name", "Role", "Responsibility"]]
            for s in stakeholders:
                tdata.append([s.get("name",""), s.get("role",""), s.get("responsibility","")])
            t = Table(tdata, colWidths=[35*mm, 40*mm, 95*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f0f0ff")]),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        # 7. Business Requirements
        reqs = brd.get("business_requirements", [])
        if reqs:
            story.append(Paragraph("7. Business Requirements", h2_style))
            tdata = [["ID", "Description", "Priority"]]
            for r in reqs:
                tdata.append([str(r.get("id","")), str(r.get("description","")), str(r.get("priority",""))])
            t = Table(tdata, colWidths=[18*mm, 130*mm, 22*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        # 8. Functional Requirements
        freqs = brd.get("functional_requirements", [])
        if freqs:
            story.append(Paragraph("8. Functional Requirements", h2_style))
            tdata = [["ID", "Description", "Priority"]]
            for r in freqs:
                tdata.append([str(r.get("id","")), str(r.get("description","")), str(r.get("priority",""))])
            t = Table(tdata, colWidths=[18*mm, 130*mm, 22*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        # 9. Non-Functional
        nfrs = brd.get("non_functional_requirements", {})
        if nfrs:
            story.append(Paragraph("9. Non-Functional Requirements", h2_style))
            for k, v in nfrs.items():
                story.append(Paragraph(f"<b>{k.replace('_', ' ').title()}</b>: {v}", body_style))
            story.append(Spacer(1, 5*mm))

        _add_list_section(story, "10. Assumptions", brd.get("assumptions", []), h2_style, body_style)
        _add_list_section(story, "11. Constraints", brd.get("constraints", []), h2_style, body_style)
        _add_list_section(story, "12. Dependencies", brd.get("dependencies", []), h2_style, body_style)

        # 13. Risks
        risks = brd.get("risks", [])
        if risks:
            story.append(Paragraph("13. Risks", h2_style))
            tdata = [["ID", "Description", "Impact", "Mitigation"]]
            for r in risks:
                tdata.append([str(r.get("id","")), str(r.get("description","")), str(r.get("impact","")), str(r.get("mitigation",""))])
            t = Table(tdata, colWidths=[13*mm, 55*mm, 18*mm, 64*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        _add_list_section(story, "14. Acceptance Criteria", brd.get("acceptance_criteria", []), h2_style, body_style)

        # 15. Timeline
        tm = brd.get("timeline_milestones", [])
        if tm:
            story.append(Paragraph("15. Timeline / Milestones", h2_style))
            tdata = [["Phase", "Description", "Target Date"]]
            for m in tm:
                tdata.append([str(m.get("phase","")), str(m.get("description","")), str(m.get("target_date",""))])
            t = Table(tdata, colWidths=[25*mm, 110*mm, 35*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), ACCENT),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.lightgrey),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 5*mm))

        story.append(PageBreak())

    # ── FRD Section ──────────────────────────────────────────
    frd = document_data.get("frd")
    if frd:
        story.append(Paragraph("FUNCTIONAL REQUIREMENTS DOCUMENT", h1_style))
        _add_section(story, "System Overview", frd.get("system_overview", {}).get("content"), h2_style, body_style)

        frs = frd.get("functional_requirements", [])
        if frs:
            story.append(Paragraph("Functional Requirements", h2_style))
            for fr in frs:
                story.append(Paragraph(f"<b>{fr.get('id','')}: {fr.get('title','')}</b> [{fr.get('priority','')}]", body_style))
                story.append(Paragraph(fr.get("description",""), body_style))

        nfrs = frd.get("non_functional_requirements", [])
        if nfrs:
            story.append(Paragraph("Non-Functional Requirements", h2_style))
            for nfr in nfrs:
                story.append(Paragraph(f"<b>{nfr.get('id','')}</b> [{nfr.get('category','')}]: {nfr.get('requirement','')}", body_style))
                if nfr.get("metric"):
                    story.append(Paragraph(f"Metric: {nfr['metric']}", meta_style))

        story.append(PageBreak())

    # ── Gap Report ───────────────────────────────────────────
    gap = document_data.get("gap")
    if gap and gap.get("gaps"):
        story.append(Paragraph("GAP ANALYSIS REPORT", h1_style))
        for g in gap["gaps"]:
            sev = g.get("severity", "LOW")
            sev_hex = {"HIGH": "cc0000", "MEDIUM": "e07000", "LOW": "1a7a1a"}.get(sev, "555555")
            story.append(Paragraph(
                f"<font color='#{sev_hex}'><b>[{sev}]</b></font> {g.get('type','').replace('_',' ').title()}",
                body_style
            ))
            story.append(Paragraph(g.get("description",""), body_style))
            story.append(Paragraph(f"<i>Recommendation: {g.get('recommendation','')}</i>", meta_style))
            story.append(Spacer(1, 3*mm))

    # ── Risk Report ──────────────────────────────────────────
    risk = document_data.get("risk")
    if risk and risk.get("risks"):
        story.append(PageBreak())
        story.append(Paragraph("RISK REGISTER", h1_style))
        for r in risk["risks"]:
            story.append(Paragraph(f"<b>{r.get('title','')} [{r.get('category','')}]</b>", body_style))
            story.append(Paragraph(f"Risk Score: <b>{r.get('risk_score','')}</b> | P: {r.get('probability','')} | I: {r.get('impact','')}", meta_style))
            story.append(Paragraph(r.get("description",""), body_style))
            story.append(Paragraph(f"Mitigation: {r.get('mitigation_strategy','')}", meta_style))
            story.append(Spacer(1, 3*mm))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _add_section(story, title, content, h_style, b_style):
    if content and content != "INSUFFICIENT_DATA":
        story.append(Paragraph(title, h_style))
        story.append(Paragraph(content.replace("\n", "<br/>"), b_style))


def _add_list_section(story, title, items, h_style, b_style):
    if items:
        story.append(Paragraph(title, h_style))
        for item in items:
            story.append(Paragraph(f"• {item}", b_style))


# ── DOCX Export (python-docx) ─────────────────────────────────
def generate_docx(document_data: dict) -> bytes:
    """Generate a .docx file from document data. Returns bytes."""
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # Title
    title = doc.add_heading(document_data.get("title", "Requirements Document"), 0)
    title.runs[0].font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)

    doc.add_paragraph(
        f"Domain: {document_data.get('domain','generic').upper()} | "
        f"Generated: {datetime.now().strftime('%B %d, %Y')} | Powered by RequirementIQ"
    ).runs[0].font.size = Pt(9)

    # BRD
    brd = document_data.get("brd")
    if brd:
        doc.add_heading("Business Requirements Document", 1)
        _docx_section(doc, "Executive Summary", brd.get("executive_summary", {}).get("content"))
        _docx_section(doc, "Problem Statement", brd.get("problem_statement", {}).get("content"))
        _docx_list(doc, "Business Objectives", brd.get("business_objectives", []))
        _docx_list(doc, "Scope — In", brd.get("scope_in", []))
        _docx_list(doc, "Scope — Out", brd.get("scope_out", []))
        _docx_list(doc, "Assumptions", brd.get("assumptions", []))
        _docx_list(doc, "Constraints", brd.get("constraints", []))
        _docx_list(doc, "Success Metrics", brd.get("success_metrics", []))

    # FRD
    frd = document_data.get("frd")
    if frd:
        doc.add_page_break()
        doc.add_heading("Functional Requirements Document", 1)
        _docx_section(doc, "System Overview", frd.get("system_overview", {}).get("content"))
        for fr in frd.get("functional_requirements", []):
            doc.add_heading(f"{fr.get('id')}: {fr.get('title')} [{fr.get('priority')}]", 3)
            doc.add_paragraph(fr.get("description", ""))

    # Gap
    gap = document_data.get("gap")
    if gap and gap.get("gaps"):
        doc.add_page_break()
        doc.add_heading("Gap Analysis Report", 1)
        for g in gap["gaps"]:
            doc.add_heading(f"[{g.get('severity')}] {g.get('type','').replace('_',' ').title()}", 3)
            doc.add_paragraph(g.get("description", ""))
            p = doc.add_paragraph()
            r = p.add_run(f"Recommendation: {g.get('recommendation','')}")
            r.italic = True

    # Risk
    risk = document_data.get("risk")
    if risk and risk.get("risks"):
        doc.add_page_break()
        doc.add_heading("Risk Register", 1)
        for r in risk["risks"]:
            doc.add_heading(f"{r.get('title')} [{r.get('category')}] — {r.get('risk_score')}", 3)
            doc.add_paragraph(r.get("description", ""))
            p = doc.add_paragraph()
            run = p.add_run(f"Mitigation: {r.get('mitigation_strategy','')}")
            run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _docx_section(doc, title, content):
    if content and content != "INSUFFICIENT_DATA":
        doc.add_heading(title, 2)
        doc.add_paragraph(content)


def _docx_list(doc, title, items):
    if items:
        doc.add_heading(title, 2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
